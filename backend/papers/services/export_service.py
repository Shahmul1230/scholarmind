import re
from io import BytesIO

from django.utils.timezone import localtime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from ..models import ChatMessage, DocumentAnalysis


def safe_export_filename(title, suffix):
    clean_title = re.sub(r"[^A-Za-z0-9_-]+", "_", title or "document")
    clean_title = clean_title.strip("_")[:80] or "document"
    return f"{clean_title}_{suffix}.docx"


def clean_inline_markdown(text):
    if not text:
        return ""

    text = str(text)

    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("`", "")
    text = text.replace("###", "")
    text = text.replace("##", "")
    text = text.replace("#", "")

    return text.strip()


def truncate_text(text, max_chars=1800):
    if not text:
        return ""

    text = str(text).strip()

    if len(text) <= max_chars:
        return text

    return text[:max_chars].rsplit(" ", 1)[0].strip() + "..."


def setup_document_styles(document):
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)


def add_export_title(document, title, subtitle=None):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(20)

    if subtitle:
        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(11)

    document.add_paragraph()


def add_key_value_table(document, items):
    valid_items = [(k, v) for k, v in items if v not in [None, ""]]

    if not valid_items:
        return

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for key, value in valid_items:
        row = table.add_row()
        row.cells[0].text = str(key)
        row.cells[1].text = str(value)

        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    document.add_paragraph()


def add_text_block(document, text):
    if not text:
        document.add_paragraph("No content available.")
        return

    lines = str(text).splitlines()

    for raw_line in lines:
        line = raw_line.rstrip()

        if not line.strip():
            continue

        stripped = line.strip()

        if stripped.startswith("### "):
            document.add_heading(clean_inline_markdown(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            document.add_heading(clean_inline_markdown(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            document.add_heading(clean_inline_markdown(stripped[2:]), level=1)
        elif re.match(r"^\s*[-*]\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(clean_inline_markdown(re.sub(r"^\s*[-*]\s+", "", stripped)))
        elif re.match(r"^\s*\d+\.\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(clean_inline_markdown(re.sub(r"^\s*\d+\.\s+", "", stripped)))
        else:
            document.add_paragraph(clean_inline_markdown(stripped))


def format_score(source):
    score = source.get("relevance_score")

    if score is None:
        return "N/A"

    try:
        return f"{float(score) * 100:.0f}%"
    except Exception:
        return str(score)


def add_sources(document, sources):
    if not sources:
        return

    document.add_heading("Sources", level=3)

    for index, source in enumerate(sources, start=1):
        source_id = source.get("source_id") or f"Source {index}"
        section_title = source.get("section_title") or "Unknown"
        page_number = source.get("page_number") or "Unknown"
        start_line = source.get("start_line") or ""
        end_line = source.get("end_line") or ""
        relevance_label = source.get("relevance_label") or "Source"
        relevance_score = format_score(source)

        document.add_heading(f"{source_id} — Page {page_number}", level=4)

        line_range = ""
        if start_line and end_line:
            line_range = f"{start_line}-{end_line}"

        add_key_value_table(
            document,
            [
                ("Source ID", source_id),
                ("Page", page_number),
                ("Line Range", line_range),
                ("Section", section_title),
                ("Relevance", relevance_label),
                ("Score", relevance_score),
            ],
        )

        content = source.get("content") or source.get("preview") or ""
        add_text_block(document, truncate_text(content, 1400))


def build_base_document(paper, export_title):
    document = Document()
    setup_document_styles(document)

    generated_at = localtime().strftime("%Y-%m-%d %I:%M %p")

    add_export_title(
        document,
        "ScholarMind Export",
        export_title,
    )

    analysis_status = "not_created"
    document_type = "Unknown"

    if hasattr(paper, "analysis"):
        analysis_status = paper.analysis.status
        document_type = paper.analysis.document_type

    add_key_value_table(
        document,
        [
            ("Document Title", paper.title),
            ("Document Type", document_type),
            ("Document Status", paper.status),
            ("Analysis Status", analysis_status),
            ("Total Pages", paper.total_pages),
            ("Generated At", generated_at),
        ],
    )

    return document


def build_chat_history_docx(paper):
    document = build_base_document(
        paper=paper,
        export_title="Chat History Export",
    )

    messages = ChatMessage.objects.filter(paper=paper).order_by("created_at")

    document.add_heading("Chat History", level=1)

    if not messages.exists():
        document.add_paragraph("No chat history found for this document.")
    else:
        for index, message in enumerate(messages, start=1):
            role = "User" if message.role == "user" else "Assistant"
            created_at = localtime(message.created_at).strftime("%Y-%m-%d %I:%M %p")

            document.add_heading(f"{index}. {role} — {created_at}", level=2)
            add_text_block(document, message.message)

            if message.role == "assistant" and message.sources:
                add_sources(document, message.sources)

            document.add_paragraph()

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def add_dict_as_section(document, title, data):
    if not data:
        return

    document.add_heading(title, level=2)

    if isinstance(data, dict):
        for key, value in data.items():
            document.add_heading(str(key), level=3)

            if isinstance(value, (dict, list)):
                add_text_block(document, str(value))
            else:
                add_text_block(document, value)
    elif isinstance(data, list):
        for item in data:
            add_text_block(document, f"- {item}")
    else:
        add_text_block(document, str(data))


def build_analysis_docx(paper):
    try:
        analysis = DocumentAnalysis.objects.get(paper=paper)
    except DocumentAnalysis.DoesNotExist:
        raise ValueError("Document intelligence profile not found for this paper.")

    document = build_base_document(
        paper=paper,
        export_title="Document Intelligence Export",
    )

    document.add_heading("Document Intelligence Profile", level=1)

    add_key_value_table(
        document,
        [
            ("Document Type", analysis.document_type),
            ("Status", analysis.status),
            ("Created At", localtime(analysis.created_at).strftime("%Y-%m-%d %I:%M %p")),
            ("Updated At", localtime(analysis.updated_at).strftime("%Y-%m-%d %I:%M %p")),
        ],
    )

    document.add_heading("Whole Document Profile", level=2)
    add_text_block(document, analysis.profile or "No profile text available.")

    add_dict_as_section(
        document,
        "Section Summaries",
        analysis.section_summaries or {},
    )

    add_dict_as_section(
        document,
        "Key Points",
        analysis.key_points or {},
    )

    add_dict_as_section(
        document,
        "Source Coverage",
        analysis.source_coverage or {},
    )

    if analysis.error_message:
        document.add_heading("Analysis Error Message", level=2)
        add_text_block(document, analysis.error_message)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()